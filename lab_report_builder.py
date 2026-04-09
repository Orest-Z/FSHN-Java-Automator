"""
=============================================================
  Java Lab Report Builder  —  by Claude for Orest Zogju
  v4.0 — Folder-Based Screenshot Workflow
=============================================================
NEW in v4:
  • Run your Java app yourself → take as many screenshots as
    you want → save them into a folder → the script picks
    them up automatically, ordered by filename.
  • Multi-screenshot per exercise: screenshot_1.png,
    screenshot_2.png, screenshot_3.png … all get embedded.
  • Side-by-side (PDF style) OR stacked layout per exercise.
  • Java compile + run helper still available as a launcher
    so you don't have to open a terminal manually.
  • Streamlit UI with a clear 3-step workflow.

SCREENSHOT FOLDER CONVENTION
─────────────────────────────
  screenshots/
    Ushtrimi1/          ← one sub-folder per exercise
      01_default.png    ← any filenames; sorted alphabetically
      02_uppercase.png
      03_lowercase.png
    Ushtrimi2/
      01_dialog.png
      02_yes_result.png
    ...

  OR flat folder with prefixes:
    screenshots/
      1_01.png   ← exercise 1, screenshot 1
      1_02.png   ← exercise 1, screenshot 2
      2_01.png   ← exercise 2, screenshot 1

Requirements:
    pip install python-docx Pillow streamlit

Run UI:
    streamlit run lab_report_builder.py
=============================================================
"""

import os
import io
import sys
import time
import subprocess
import tempfile
import argparse
from pathlib import Path


def is_running_in_streamlit() -> bool:
    try:
        from streamlit.runtime.scriptrunner import get_script_run_ctx
        return get_script_run_ctx() is not None
    except Exception:
        return False


# ─── TEMPLATE ─────────────────────────────────────────────────────────────────
TEMPLATE = {
    "page_width_inches":   8.27,
    "page_height_inches":  11.69,
    "margin_inches":       1.0,

    "header_font":         "Times New Roman",
    "header_size_pt":      14,
    "header_bold":         True,

    "exercise_title_font": "Times New Roman",
    "exercise_title_size": 12,
    "exercise_title_bold": True,

    "description_font":    "Times New Roman",
    "description_size":    11,

    "code_font":           "Courier New",
    "code_size":           8,
    "code_bg_color":       "EEEEEE",

    # "side-by-side" = code left + screenshots right (PDF style)
    # "stacked"      = code on top, screenshots below
    "layout_mode":         "side-by-side",

    # Fraction of content width for each column (must sum to 1.0)
    "code_col_fraction":   0.58,
    "img_col_fraction":    0.42,

    # Vertical alignment of images in the right column: top / middle / bottom
    "image_valign":        "top",

    # Max image width for stacked mode
    "image_max_width_inches": 5.5,

    "space_after_code_pt":  12,
    "space_after_image_pt": 6,

    "page_break_between":  True,
}
# ──────────────────────────────────────────────────────────────────────────────


# ╔══════════════════════════════════════════════════════════════════╗
# ║            SCREENSHOT FOLDER  UTILITIES                         ║
# ╚══════════════════════════════════════════════════════════════════╝

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"}


def collect_images_from_folder(folder: str) -> list:
    """
    Collect all image files from a folder, sorted by filename.
    Returns list of {"path": abs_path, "bytes": None}.
    """
    folder = Path(folder)
    if not folder.is_dir():
        return []
    files = sorted(
        p for p in folder.iterdir()
        if p.suffix.lower() in IMAGE_EXTS and p.is_file()
    )
    return [{"path": str(f), "bytes": None} for f in files]


def collect_images_from_subfolders(root: str) -> dict:
    """
    Scan root for sub-folders. Each sub-folder name is an exercise key.
    Returns {folder_name: [image_dicts]}.

    Example structure:
      root/
        Ushtrimi1/  -> images for exercise 1
        Ushtrimi2/  -> images for exercise 2
    """
    root = Path(root)
    result = {}
    if not root.is_dir():
        return result
    for sub in sorted(root.iterdir()):
        if sub.is_dir():
            imgs = collect_images_from_folder(str(sub))
            if imgs:
                result[sub.name] = imgs
    return result


def launch_java(java_file: str) -> tuple:
    """
    Compile and run a .java file.
    Returns (success: bool, message: str, process_handle_or_None).
    """
    java_file = Path(java_file)
    if not java_file.exists():
        return False, f"File not found: {java_file}", None

    class_dir  = str(java_file.parent)
    class_name = java_file.stem

    # Compile
    result = subprocess.run(
        ["javac", str(java_file)],
        capture_output=True, text=True, cwd=class_dir,
    )
    if result.returncode != 0:
        return False, f"Compile error:\n{result.stderr}", None

    # Run (non-blocking)
    proc = subprocess.Popen(
        ["java", class_name],
        cwd=class_dir,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    return True, f"Running {class_name}", proc


def read_java_source(java_file: str) -> str:
    """Read and return the source code of a .java file."""
    try:
        with open(java_file, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    except Exception as e:
        return f"// Could not read file: {e}"


# ╔══════════════════════════════════════════════════════════════════╗
# ║                    DOCX  CORE  BUILDER                          ║
# ╚══════════════════════════════════════════════════════════════════╝

def _inches(val):
    from docx.shared import Inches
    return Inches(val)

def _pt(val):
    from docx.shared import Pt
    return Pt(val)


def _set_shading(paragraph, hex_color: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def _add_page_break(doc):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p   = doc.add_paragraph()
    run = p.add_run()
    br  = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _resize_image(source, max_width_inches: float) -> tuple:
    """Return (w_in, h_in) preserving aspect ratio within max_width_inches."""
    from PIL import Image
    if isinstance(source, (bytes, bytearray)):
        img = Image.open(io.BytesIO(source))
    else:
        img = Image.open(source)

    w_px, h_px = img.size
    dpi = img.info.get("dpi", (96, 96))
    dpi_x = float(dpi[0]) if isinstance(dpi, (tuple, list)) else 96.0
    if dpi_x <= 0:
        dpi_x = 96.0

    w_in = w_px / dpi_x
    h_in = h_px / dpi_x

    if w_in > max_width_inches:
        scale = max_width_inches / w_in
        w_in  = max_width_inches
        h_in  = h_in * scale

    return w_in, h_in


# ── Invisible table borders ────────────────────────────────────────

def _make_invisible_borders_xml():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "none")
        b.set(qn("w:sz"),    "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tc_borders.append(b)
    return tc_borders


def _set_cell_invisible(cell):
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    ns   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for old in tcPr.findall(f"{{{ns}}}tcBorders"):
        tcPr.remove(old)
    tcPr.append(_make_invisible_borders_xml())


def _set_cell_valign(cell, valign: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    map_ = {"top": "top", "middle": "center", "center": "center", "bottom": "bottom"}
    tcPr   = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), map_.get(valign, "top"))
    tcPr.append(vAlign)


# ── Fill cells ─────────────────────────────────────────────────────

def _fill_code_cell(cell, code: str, T: dict):
    """Write shaded code lines into a table cell."""
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)

    for line in code.split("\n"):
        p = cell.add_paragraph()
        _set_shading(p, T["code_bg_color"])
        p.paragraph_format.space_before = _pt(0)
        p.paragraph_format.space_after  = _pt(0)
        p.paragraph_format.left_indent  = _inches(0.05)
        run = p.add_run(line if line else " ")
        run.font.name = T["code_font"]
        run.font.size = _pt(T["code_size"])


def _fill_image_cell(cell, images: list, T: dict,
                     col_width_inches: float, valign: str = "top"):
    """Stack all images vertically inside a table cell."""
    _set_cell_valign(cell, valign)

    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)

    for img in images:
        img_path  = img.get("path")
        img_bytes = img.get("bytes")
        if not img_path and not img_bytes:
            continue
        try:
            source = img_bytes if img_bytes else img_path
            w_in, _ = _resize_image(source, col_width_inches)
            p   = cell.add_paragraph()
            p.paragraph_format.space_before = _pt(2)
            p.paragraph_format.space_after  = _pt(T["space_after_image_pt"])
            run = p.add_run()
            if img_bytes:
                run.add_picture(io.BytesIO(img_bytes), width=_inches(w_in))
            else:
                run.add_picture(img_path, width=_inches(w_in))
        except Exception as e:
            cell.add_paragraph(f"[Image error: {e}]")


# ── Layout renderers ───────────────────────────────────────────────

def _render_side_by_side(doc, ex: dict, T: dict):
    """Code LEFT column, screenshots RIGHT column — matches PDF style."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement as OE

    content_w = T["page_width_inches"] - 2 * T["margin_inches"]
    code_w    = content_w * T["code_col_fraction"]
    img_w     = content_w * T["img_col_fraction"]

    code   = ex.get("code", "")
    images = ex.get("images", [])
    valign = ex.get("image_valign", T.get("image_valign", "top"))

    def _twips(inches):
        return str(int(inches * 1440))

    # Build invisible 2-column table
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    tbl   = table._tbl
    tblPr = tbl.find(
        f'{{{tbl.nsmap.get("w","http://schemas.openxmlformats.org/wordprocessingml/2006/main")}}}tblPr'
    )

    # Total width
    tblW = OE("w:tblW")
    tblW.set(qn("w:w"), _twips(content_w))
    tblW.set(qn("w:type"), "dxa")
    if tblPr is not None:
        tblPr.append(tblW)

    # Invisible table borders
    tblBorders = OE("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OE(f"w:{side}")
        b.set(qn("w:val"),   "none")
        b.set(qn("w:sz"),    "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tblBorders.append(b)
    if tblPr is not None:
        tblPr.append(tblBorders)

    cells = table.rows[0].cells
    for cell, w_in in zip(cells, [code_w, img_w]):
        tcPr = cell._tc.get_or_add_tcPr()
        tcW  = OE("w:tcW")
        tcW.set(qn("w:w"), _twips(w_in))
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)
        _set_cell_invisible(cell)

    # Fill
    if code.strip():
        _fill_code_cell(cells[0], code, T)
    else:
        cells[0].paragraphs[0].add_run(" ")

    if images:
        _fill_image_cell(cells[1], images, T,
                         col_width_inches=img_w - 0.1,
                         valign=valign)
    else:
        cells[1].paragraphs[0].add_run(" ")

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = _pt(T["space_after_code_pt"])


def _render_stacked(doc, ex: dict, T: dict):
    """Code on top, all screenshots below — original style."""
    code   = ex.get("code", "")
    images = ex.get("images", [])

    if code.strip():
        for line in code.split("\n"):
            p = doc.add_paragraph()
            _set_shading(p, T["code_bg_color"])
            p.paragraph_format.space_before = _pt(0)
            p.paragraph_format.space_after  = _pt(0)
            p.paragraph_format.left_indent  = _inches(0.1)
            run = p.add_run(line if line else " ")
            run.font.name = T["code_font"]
            run.font.size = _pt(T["code_size"])
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = _pt(T["space_after_code_pt"])

    for img in images:
        img_path  = img.get("path")
        img_bytes = img.get("bytes")
        if not img_path and not img_bytes:
            continue
        try:
            source = img_bytes if img_bytes else img_path
            w_in, _ = _resize_image(source, T["image_max_width_inches"])
            p   = doc.add_paragraph()
            run = p.add_run()
            if img_bytes:
                run.add_picture(io.BytesIO(img_bytes), width=_inches(w_in))
            else:
                run.add_picture(img_path, width=_inches(w_in))
            p.paragraph_format.space_after = _pt(T["space_after_image_pt"])
        except Exception as e:
            doc.add_paragraph(f"[Image error: {e}]")


def _add_exercise(doc, ex: dict, T: dict):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    number      = ex.get("number", "?")
    description = ex.get("description", "")
    layout      = ex.get("layout", T.get("layout_mode", "side-by-side"))

    # Title
    tp = doc.add_paragraph()
    r  = tp.add_run(f"Ushtrimi {number}:")
    r.font.name = T["exercise_title_font"]
    r.font.size = _pt(T["exercise_title_size"])
    r.font.bold = T["exercise_title_bold"]
    tp.paragraph_format.space_before = _pt(10)
    tp.paragraph_format.space_after  = _pt(6)

    # Description
    if description.strip():
        for line in description.strip().split("\n"):
            p   = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = T["description_font"]
            run.font.size = _pt(T["description_size"])
            p.paragraph_format.space_after = _pt(4)

    doc.add_paragraph()  # gap before code block

    if layout == "side-by-side":
        _render_side_by_side(doc, ex, T)
    else:
        _render_stacked(doc, ex, T)


# ── Document assembly ──────────────────────────────────────────────

def create_document(
    doc_title:   str,
    author_name: str,
    exercises:   list,
    output_path: str  = "lab_report.docx",
    template:    dict = None,
) -> str:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    T   = template or TEMPLATE
    doc = Document()

    # Page setup
    sec = doc.sections[0]
    sec.page_width   = _inches(T["page_width_inches"])
    sec.page_height  = _inches(T["page_height_inches"])
    m = _inches(T["margin_inches"])
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = m

    # Normal style defaults
    sty = doc.styles["Normal"]
    sty.font.name = T["description_font"]
    sty.font.size = _pt(T["description_size"])
    sty.paragraph_format.space_before = _pt(0)
    sty.paragraph_format.space_after  = _pt(4)

    # Header: "Title    [tab]    Author"
    hp = doc.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = hp.add_run(doc_title)
    r1.font.name = T["header_font"]
    r1.font.size = _pt(T["header_size_pt"])
    r1.font.bold = T["header_bold"]
    hp.add_run("\t")
    r2 = hp.add_run(author_name)
    r2.font.name = T["header_font"]
    r2.font.size = _pt(T["header_size_pt"])
    r2.font.bold = T["header_bold"]

    # Right tab stop
    pPr      = hp._p.get_or_add_pPr()
    tabs_el  = OxmlElement("w:tabs")
    tab_el   = OxmlElement("w:tab")
    cw_twips = int((T["page_width_inches"] - 2 * T["margin_inches"]) * 1440)
    tab_el.set(qn("w:val"), "right")
    tab_el.set(qn("w:pos"), str(cw_twips))
    tabs_el.append(tab_el)
    pPr.append(tabs_el)

    # Horizontal rule
    bp   = doc.add_paragraph()
    bPr  = bp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "000000")
    pBdr.append(bot)
    bPr.append(pBdr)
    bp.paragraph_format.space_after = _pt(10)

    # Exercises
    for idx, ex in enumerate(exercises):
        _add_exercise(doc, ex, T)
        if T["page_break_between"] and idx < len(exercises) - 1:
            _add_page_break(doc)

    doc.save(output_path)
    return output_path


# ╔══════════════════════════════════════════════════════════════════╗
# ║                       STREAMLIT  UI                             ║
# ╚══════════════════════════════════════════════════════════════════╝

def run_streamlit():
    import streamlit as st

    st.set_page_config(
        page_title = "Java Lab Report Builder",
        page_icon  = "☕",
        layout     = "wide",
    )

    st.title("☕ Java Lab Report Builder  v4")

    # ── Sidebar settings ─────────────────────────────────────────────
    with st.sidebar:
        st.header("⚙️ Document Settings")
        doc_title   = st.text_input("Assignment Title", "Laborator Java 2")
        author_name = st.text_input("Name / Group",     "Orest Zogju GR B1.2")
        output_name = st.text_input("Output filename",  "lab_report.docx")

        st.divider()
        st.subheader("📐 Layout & Fonts")
        img_width   = st.slider("Image width — stacked mode (inches)", 2.0, 8.0, 5.5, 0.5)
        code_size   = st.slider("Code font size (pt)", 6, 14, 8)
        desc_size   = st.slider("Description font size (pt)", 8, 14, 11)
        code_frac   = st.slider("Code column width — side-by-side (%)", 40, 80, 58) / 100
        img_frac    = round(1.0 - code_frac, 2)
        st.caption(f"↳ Image column: {int(img_frac * 100)}%")
        page_breaks = st.checkbox("Page break between exercises", value=True)

    def _make_template():
        return {
            **TEMPLATE,
            "image_max_width_inches": img_width,
            "code_size":              code_size,
            "description_size":       desc_size,
            "code_col_fraction":      code_frac,
            "img_col_fraction":       img_frac,
            "page_break_between":     page_breaks,
        }

    # ── Session state ─────────────────────────────────────────────────
    if "exercises"  not in st.session_state:
        st.session_state.exercises  = []
    if "java_proc"  not in st.session_state:
        st.session_state.java_proc  = None
    if "java_file"  not in st.session_state:
        st.session_state.java_file  = ""

    # ─────────────────────────────────────────────────────────────────
    # STEP 1 — Java Launcher
    # ─────────────────────────────────────────────────────────────────
    st.header("① Launch Java Exercise")
    st.caption(
        "Compile and run a .java file so you can interact with it. "
        "Take screenshots manually using any tool (Snipping Tool, PrintScreen, etc.) "
        "and save them into a folder."
    )

    col_file, col_btn = st.columns([3, 1])
    with col_file:
        java_path_input = st.text_input(
            "Path to .java file",
            placeholder="C:/Users/Orest/JavaLab/Ushtrimi1.java",
            label_visibility="collapsed",
        )
    with col_btn:
        launch_btn = st.button("▶ Compile & Run", use_container_width=True)

    if launch_btn:
        if not java_path_input or not os.path.isfile(java_path_input):
            st.error("Enter a valid path to a .java file.")
        else:
            # Kill previous process
            if st.session_state.java_proc:
                try:
                    st.session_state.java_proc.terminate()
                except Exception:
                    pass

            ok, msg, proc = launch_java(java_path_input)
            if ok:
                st.session_state.java_proc = proc
                st.session_state.java_file = java_path_input
                st.success(
                    f"✅ {msg}  — The Swing window should appear now.  "
                    "Interact with it, take your screenshots, save them to a folder, "
                    "then continue to Step ②."
                )
            else:
                st.error(msg)
                st.session_state.java_proc = None

    # Process status
    if st.session_state.java_proc:
        if st.session_state.java_proc.poll() is not None:
            st.info("ℹ️ The Java process has already exited.")
            st.session_state.java_proc = None
        else:
            st.success(
                f"🟢 Java is running: `{Path(st.session_state.java_file).name}`"
            )
            if st.button("⏹ Stop Java process"):
                st.session_state.java_proc.terminate()
                st.session_state.java_proc = None
                st.success("Process stopped.")

    st.divider()

    # ─────────────────────────────────────────────────────────────────
    # STEP 2 — Build Exercise Entry
    # ─────────────────────────────────────────────────────────────────
    st.header("② Add Exercise to Document")

    c1, c2 = st.columns([1, 2])

    with c1:
        ex_number = st.text_input(
            "Exercise Number",
            value=str(len(st.session_state.exercises) + 1),
        )
        layout_choice = st.selectbox(
            "Image Layout",
            [
                "Side-by-Side  (code left, screenshots right — PDF style)",
                "Stacked  (screenshots below code)",
            ],
        )
        valign_choice = st.selectbox(
            "Screenshot Vertical Align  (side-by-side only)",
            ["Top", "Middle", "Bottom"],
        )

        st.markdown("---")
        st.markdown("**📁 Load screenshots from folder**")
        st.caption(
            "Save all screenshots for this exercise into one folder "
            "(e.g. `C:/Screenshots/Ushtrimi1/`). "
            "Files are loaded in alphabetical order — name them "
            "`01_empty.png`, `02_valid.png`, `03_error.png` etc. "
            "to control the order."
        )
        img_folder = st.text_input(
            "Screenshot folder",
            placeholder="C:/Screenshots/Ushtrimi1",
            label_visibility="collapsed",
            key="img_folder_input",
        )

        st.markdown("**— or upload directly —**")
        uploaded_files = st.file_uploader(
            "Upload screenshots  (multi-select OK)",
            type=["png", "jpg", "jpeg"],
            accept_multiple_files=True,
            key="screenshot_uploader",
        )

    with c2:
        # Auto-fill code from the currently launched java file
        default_code = ""
        if st.session_state.java_file and os.path.isfile(st.session_state.java_file):
            default_code = read_java_source(st.session_state.java_file)

        description = st.text_area(
            "Exercise Description",
            height=110,
            placeholder="Ndertoni nje dritare me titull…",
        )
        code = st.text_area(
            "Java Code  (auto-filled from the launched file above)",
            value=default_code,
            height=280,
        )

    # ── Resolve images ────────────────────────────────────────────────
    images_from_folder  = []
    images_from_uploads = []

    if img_folder:
        if os.path.isdir(img_folder):
            images_from_folder = collect_images_from_folder(img_folder)
            if images_from_folder:
                st.success(
                    f"📂 Found **{len(images_from_folder)}** image(s) in "
                    f"`{img_folder}`  (sorted by filename)"
                )
                with st.expander("Preview folder images", expanded=False):
                    thumb_cols = st.columns(min(len(images_from_folder), 5))
                    for i, img in enumerate(images_from_folder):
                        thumb_cols[i % 5].image(
                            img["path"],
                            caption=Path(img["path"]).name,
                            width=150,
                        )
            else:
                st.warning(
                    f"Folder `{img_folder}` exists but contains no supported image files."
                )
        else:
            st.error(f"Folder not found: `{img_folder}`")

    if uploaded_files:
        images_from_uploads = [{"path": None, "bytes": f.read()} for f in uploaded_files]
        st.success(f"⬆️  {len(images_from_uploads)} uploaded screenshot(s) ready.")

    all_images = images_from_folder + images_from_uploads

    if all_images:
        st.info(
            f"**Total screenshots for Ushtrimi {ex_number}: {len(all_images)}**  "
            f"({len(images_from_folder)} from folder + {len(images_from_uploads)} uploaded)"
        )

    # ── Add button ────────────────────────────────────────────────────
    if st.button("➕ Add Exercise to Document", type="primary"):
        if not ex_number.strip():
            st.error("Enter an exercise number.")
        elif not code.strip() and not all_images:
            st.warning("Add at least some code or a screenshot before adding.")
        else:
            layout_val = "side-by-side" if "Side" in layout_choice else "stacked"
            st.session_state.exercises.append({
                "number":       ex_number.strip(),
                "description":  description,
                "code":         code,
                "images":       all_images,
                "layout":       layout_val,
                "image_valign": valign_choice.lower(),
            })
            st.success(
                f"✅ Ushtrimi {ex_number} added — "
                f"{len(code.splitlines())} code lines, "
                f"{len(all_images)} screenshot(s), "
                f"layout: {layout_val}"
            )

    st.divider()

    # ─────────────────────────────────────────────────────────────────
    # STEP 3 — Review Queue & Generate
    # ─────────────────────────────────────────────────────────────────
    st.header("③ Review Queue & Generate .docx")

    if not st.session_state.exercises:
        st.info("No exercises added yet. Use Step ② above.")
    else:
        st.subheader(f"📋 Queue — {len(st.session_state.exercises)} exercise(s)")

        for i, ex in enumerate(st.session_state.exercises):
            n_img = len(ex.get("images", []))
            with st.expander(
                f"Ushtrimi {ex['number']}  │  "
                f"{len(ex['code'].splitlines())} code lines  │  "
                f"🖼️ ×{n_img}  │  layout: {ex.get('layout','?')}",
                expanded=False,
            ):
                c_edit, c_del = st.columns([5, 1])
                with c_edit:
                    new_desc = st.text_area(
                        "Description",
                        value=ex["description"],
                        key=f"desc_{i}",
                        height=80,
                    )
                    st.session_state.exercises[i]["description"] = new_desc

                    new_layout = st.selectbox(
                        "Layout",
                        ["side-by-side", "stacked"],
                        index=0 if ex.get("layout") == "side-by-side" else 1,
                        key=f"layout_{i}",
                    )
                    st.session_state.exercises[i]["layout"] = new_layout

                    # Screenshot thumbnails
                    imgs = ex.get("images", [])
                    if imgs:
                        st.markdown(f"**Screenshots ({len(imgs)}):**")
                        thumb_cols = st.columns(min(len(imgs), 5))
                        for j, img in enumerate(imgs):
                            try:
                                if img.get("bytes"):
                                    thumb_cols[j % 5].image(
                                        img["bytes"],
                                        caption=f"#{j+1}",
                                        width=120,
                                    )
                                elif img.get("path"):
                                    thumb_cols[j % 5].image(
                                        img["path"],
                                        caption=Path(img["path"]).name,
                                        width=120,
                                    )
                            except Exception:
                                thumb_cols[j % 5].caption(f"[img {j+1} error]")

                with c_del:
                    if st.button("🗑️ Remove", key=f"del_{i}"):
                        st.session_state.exercises.pop(i)
                        st.rerun()

        st.divider()
        col_clear, col_gen = st.columns(2)

        with col_clear:
            if st.button("🗑️ Clear All Exercises"):
                st.session_state.exercises = []
                st.rerun()

        with col_gen:
            if st.button("📥 Generate .docx", type="primary", use_container_width=True):
                with st.spinner("Building document…"):
                    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                        tmp_path = tmp.name

                    create_document(
                        doc_title   = doc_title,
                        author_name = author_name,
                        exercises   = st.session_state.exercises,
                        output_path = tmp_path,
                        template    = _make_template(),
                    )

                    buf = io.BytesIO()
                    with open(tmp_path, "rb") as f:
                        buf.write(f.read())
                    os.unlink(tmp_path)
                    buf.seek(0)

                st.download_button(
                    label     = f"⬇️  Download  {output_name}",
                    data      = buf,
                    file_name = output_name,
                    mime      = "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
                st.success("✅ Document ready — click the button above to save it.")

    # ── Batch import helper ──────────────────────────────────────────
    st.divider()
    with st.expander("⚡ Batch import — load ALL exercises from sub-folders at once"):
        st.markdown(
            """
            **Expected folder structure:**
            ```
            root_folder/
              Ushtrimi1/   ← screenshots for exercise 1
              Ushtrimi2/   ← screenshots for exercise 2
              Ushtrimi3/   ← ...
            ```
            This creates exercise entries for each sub-folder automatically.
            You can then edit descriptions and code in Step ③.
            """
        )
        batch_root = st.text_input(
            "Root screenshots folder",
            placeholder="C:/Screenshots/Lab2",
            key="batch_root",
        )
        if st.button("📂 Import all sub-folders"):
            if not batch_root or not os.path.isdir(batch_root):
                st.error("Enter a valid folder path.")
            else:
                sub_map = collect_images_from_subfolders(batch_root)
                if not sub_map:
                    st.warning("No sub-folders with images found.")
                else:
                    for folder_name, imgs in sub_map.items():
                        st.session_state.exercises.append({
                            "number":       folder_name,
                            "description":  "",
                            "code":         "",
                            "images":       imgs,
                            "layout":       "side-by-side",
                            "image_valign": "top",
                        })
                    st.success(
                        f"✅ Imported {len(sub_map)} exercise(s) from sub-folders. "
                        "Add code and descriptions in Step ③."
                    )
                    st.rerun()


# ╔══════════════════════════════════════════════════════════════════╗
# ║                      CLI  +  DEMO                               ║
# ╚══════════════════════════════════════════════════════════════════╝

def run_cli():
    print("\n╔══════════════════════════════════════╗")
    print("║   Java Lab Report Builder  (CLI)    ║")
    print("╚══════════════════════════════════════╝\n")

    doc_title   = input("Assignment title  [Laborator Java 2]: ").strip() or "Laborator Java 2"
    author_name = input("Your name/group   [Orest Zogju GR B1.2]: ").strip() or "Orest Zogju GR B1.2"
    output_path = input("Output filename   [lab_report.docx]: ").strip() or "lab_report.docx"
    layout_def  = input("Default layout    [side/stacked]: ").strip() or "side"
    layout_def  = "side-by-side" if layout_def.startswith("s") and "t" not in layout_def else "stacked"

    exercises = []
    print("\nType END to finish multi-line blocks. DONE for exercise number to finish.\n")

    while True:
        print(f"\n─── Exercise {len(exercises)+1} ───────────────────────")
        num = input("  Exercise number (or DONE): ").strip()
        if num.upper() == "DONE":
            break

        # Description
        desc_lines = []
        print("  Description (END to finish):")
        while True:
            l = input()
            if l.strip() == "END": break
            desc_lines.append(l)

        # Code: path to .java or paste
        java_file = input("  Path to .java file (blank = paste code manually): ").strip()
        proc = None
        if java_file and os.path.isfile(java_file):
            code = read_java_source(java_file)
            print(f"  ✓ Loaded {len(code.splitlines())} lines from {java_file}")
            run_it = input("  Compile and run it now? [y/N]: ").strip().lower() == "y"
            if run_it:
                ok, msg, proc = launch_java(java_file)
                print(f"  {'✓' if ok else '✗'} {msg}")
        else:
            code_lines = []
            print("  Java code (END to finish):")
            while True:
                l = input()
                if l.strip() == "END": break
                code_lines.append(l)
            code = "\n".join(code_lines)

        # Screenshots
        images = []
        img_input = input(
            "  Screenshot folder path "
            "(or blank to enter image paths one by one): "
        ).strip()
        if img_input and os.path.isdir(img_input):
            images = collect_images_from_folder(img_input)
            print(f"  ✓ Loaded {len(images)} image(s) from {img_input}")
        else:
            if img_input:
                print(f"  ⚠️  Folder not found: {img_input}")
            print("  Enter screenshot paths one by one (blank to stop):")
            while True:
                p = input("    path: ").strip()
                if not p: break
                if os.path.isfile(p):
                    images.append({"path": p, "bytes": None})
                else:
                    print(f"    ⚠️  Not found: {p}")

        # Kill Java if we launched it
        if proc:
            input("  Press ENTER when done with the Java window (it will be closed)… ")
            proc.terminate()

        exercises.append({
            "number":       num,
            "description":  "\n".join(desc_lines),
            "code":         code,
            "images":       images,
            "layout":       layout_def,
            "image_valign": "top",
        })
        print(f"  ✓ Ushtrimi {num} added — {len(images)} screenshot(s)")

    if not exercises:
        print("No exercises added. Exiting.")
        return

    path = create_document(doc_title, author_name, exercises, output_path)
    print(f"\n✅ Report saved: {os.path.abspath(path)}\n")


def run_demo(output_path="lab_report_demo_v4.docx"):
    exercises = [
        {
            "number":      "1",
            "description": (
                "Ndertoni nje dritare me titull Sistemi i Autorizimit. "
                "Dritarja duhet te kete permasa fikse, te shfaqet ne qender "
                "te ekranit dhe te kete nje label me tekstin \"Ju lutem identifikohuni\"."
            ),
            "code": (
                "import javax.swing.*;\n"
                "import java.awt.*;\n\n"
                "public class Ushtrimi1 extends JFrame {\n"
                "    Ushtrimi1() {\n"
                "        JFrame frame = new JFrame(\"Sistemi i Autorizimit\");\n"
                "        frame.setSize(400, 400);\n"
                "        frame.setLayout(new FlowLayout());\n"
                "        JLabel l1 = new JLabel(\"Ju lutem identifikohuni\");\n"
                "        JButton btn1 = new JButton(\"X\");\n"
                "        btn1.setBackground(Color.RED);\n"
                "        btn1.addActionListener(e -> System.exit(0));\n"
                "        frame.add(l1);\n"
                "        frame.add(btn1);\n"
                "        frame.setDefaultCloseOperation(JFrame.EXIT_ON_CLOSE);\n"
                "        frame.setLocationRelativeTo(null);\n"
                "        frame.setVisible(true);\n"
                "    }\n"
                "    public static void main(String[] args) { new Ushtrimi1(); }\n"
                "}"
            ),
            "images":       [],
            "layout":       "side-by-side",
            "image_valign": "top",
        },
        {
            "number":      "5",
            "description": (
                "Ndertoni nje dritare me nje fushe teksti per \"Kodin e Kuponit\" "
                "dhe nje buton \"Apliko\". Nese kodi eshte \"STUDENT2026\", "
                "label-i te thote \"Zbritja u aplikua: 20%\" me ngjyre jeshile. "
                "Per cdo kod tjeter, label-i te thote \"Kod i pasakte\" me ngjyre te kuqe."
            ),
            "code": (
                "import javax.swing.*;\n"
                "import java.awt.*;\n\n"
                "public class Ushtrimi5 {\n"
                "    Ushtrimi5() {\n"
                "        JFrame frame = new JFrame();\n"
                "        frame.setSize(400, 400);\n"
                "        frame.setLayout(new FlowLayout());\n"
                "        JLabel l1 = new JLabel(\"Kodi i kuponit: \");\n"
                "        JLabel l2 = new JLabel(\"\");\n"
                "        l2.setFont(new Font(\"Times New Roman\", Font.BOLD, 20));\n"
                "        JTextField t1 = new JTextField(8);\n"
                "        JButton b1 = new JButton(\"Apliko\");\n"
                "        b1.setBackground(Color.GREEN);\n"
                "        b1.addActionListener(e -> {\n"
                "            if (t1.getText().equals(\"STUDENT2026\")) {\n"
                "                l2.setText(\"Zbritja u aplikua me 20%\");\n"
                "                l2.setForeground(Color.GREEN);\n"
                "            } else {\n"
                "                l2.setText(\"Kod i gabuar\");\n"
                "                l2.setForeground(Color.RED);\n"
                "            }\n"
                "        });\n"
                "        frame.add(l2); frame.add(t1);\n"
                "        frame.add(l1); frame.add(b1);\n"
                "        frame.setLocationRelativeTo(null);\n"
                "        frame.setDefaultCloseOperation(JFrame.EXIT_ON_CLOSE);\n"
                "        frame.setVisible(true);\n"
                "    }\n"
                "    public static void main(String[] args) { new Ushtrimi5(); }\n"
                "}"
            ),
            "images":       [],
            "layout":       "side-by-side",
            "image_valign": "top",
        },
    ]

    path = create_document(
        "Laborator Java 2", "Orest Zogju GR B1.2",
        exercises, output_path,
    )
    print(f"Demo created: {os.path.abspath(path)}")
    return path


# ╔══════════════════════════════════════════════════════════════════╗
# ║                         ENTRY POINT                             ║
# ╚══════════════════════════════════════════════════════════════════╝

if __name__ == "__main__":
    if is_running_in_streamlit():
        run_streamlit()
    else:
        parser = argparse.ArgumentParser(description="Java Lab Report Builder v4")
        parser.add_argument("--demo", action="store_true", help="Generate a demo .docx")
        parser.add_argument("--cli",  action="store_true", help="Run interactive CLI")
        args = parser.parse_args()

        if args.demo:
            run_demo()
        elif args.cli:
            run_cli()
        else:
            run_cli()

# Streamlit auto-entry
try:
    import streamlit as _st
    if _st.runtime.exists():
        run_streamlit()
except Exception:
    pass
